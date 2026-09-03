import io
from unittest.mock import patch

from docx import Document


def test_upload_valid_docx_returns_version_one(client):
    document = Document()
    document.add_paragraph('Senior software engineer with 8 years of Python, FastAPI, and PostgreSQL experience.')
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = client.post(
        '/api/resumes?session_id=session-123',
        files={'file': ('resume.docx', buffer.read(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload['version'] == 1
    assert payload['extraction_method'] == 'python-docx'
    assert payload['extraction_status'] == 'success'
    assert payload['extraction_error'] is None
    assert payload['session_id'] == 'session-123'


def test_upload_marks_failed_extraction_when_provider_errors(client):
    document = Document()
    document.add_paragraph('Senior software engineer with 8 years of Python, FastAPI, and PostgreSQL experience.')
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    with patch('app.routers.upload.get_llm_provider', side_effect=RuntimeError('missing API key')):
        response = client.post(
            '/api/resumes?session_id=session-failed',
            files={'file': ('resume.docx', buffer.read(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')},
        )

    assert response.status_code == 200
    payload = response.json()
    assert payload['extraction_status'] == 'failed'
    assert payload['extraction_error'] == 'missing API key'
    assert payload['raw_text_preview']


def test_upload_rejects_unsupported_type(client):
    response = client.post(
        '/api/resumes',
        files={'file': ('resume.txt', b'this is not a resume', 'text/plain')},
    )
    assert response.status_code == 400
    assert 'Unsupported file type' in response.json()['detail']


def test_second_upload_in_same_session_increments_version(client):
    document = Document()
    document.add_paragraph('Senior software engineer with 8 years of Python, FastAPI, and PostgreSQL experience across multiple teams.')
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    first = client.post(
        '/api/resumes?session_id=session-456',
        files={'file': ('resume-one.docx', buffer.read(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')},
    )
    buffer.seek(0)
    second = client.post(
        '/api/resumes?session_id=session-456',
        files={'file': ('resume-two.docx', buffer.read(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')},
    )

    assert first.status_code == 200
    assert second.status_code == 200
    assert first.json()['version'] == 1
    assert second.json()['version'] == 2
