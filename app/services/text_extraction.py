from dataclasses import dataclass
from pathlib import Path

import fitz
import pdfplumber
from docx import Document

MIN_VIABLE_CHARS = 50


@dataclass
class ExtractionResult:
    text: str
    method: str


def _extract_pdf_text(path: Path) -> ExtractionResult:
    try:
        with pdfplumber.open(path) as pdf:
            chunks = []
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    chunks.append(page_text)
            text = "\n".join(chunks).strip()
        if len(text) >= MIN_VIABLE_CHARS:
            return ExtractionResult(text=text, method="pdfplumber")
    except Exception:
        pass

    try:
        doc = fitz.open(path)
        text_chunks = []
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text_chunks.append(page_text)
        text = "\n".join(text_chunks).strip()
        if len(text) >= MIN_VIABLE_CHARS:
            return ExtractionResult(text=text, method="pymupdf_fallback")
    except Exception:
        pass

    raise ValueError(
        "Could not extract readable text from this PDF. It is likely a scanned image PDF without a text layer."
    )


def _extract_docx_text(path: Path) -> ExtractionResult:
    doc = Document(path)
    chunks = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_parts = [cell.text.strip() for cell in row.cells]
            row_text = " | ".join(part for part in row_parts if part)
            if row_text:
                chunks.append(row_text)

    text = "\n".join(chunks).strip()
    if len(text) < MIN_VIABLE_CHARS:
        raise ValueError(
            "Document text is too short to qualify as a meaningful resume. Please provide a fuller resume text."
        )

    return ExtractionResult(text=text, method="python-docx")


def extract_resume_text(file_path: str | Path) -> ExtractionResult:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf_text(path)
    if suffix == ".docx":
        return _extract_docx_text(path)

    raise ValueError(f"Unsupported file type for extraction: {suffix or 'unknown'}")
